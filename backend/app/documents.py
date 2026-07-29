import re
import shutil
import time
from pathlib import Path
from fastapi import HTTPException, UploadFile
from sqlalchemy.orm import Session
from app.models import Document, Chunk
from app.config import settings
from app.embedding import embedding_engine
from app.vector_store import VectorStore

ALLOWED_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}


def process_upload(file: UploadFile, user_id: int, db: Session, vector_store: VectorStore) -> Document:
    if file.content_type not in ALLOWED_TYPES:
        raise HTTPException(400, f"Unsupported file type: {file.content_type}")

    upload_dir = Path(settings.UPLOAD_DIR)
    upload_dir.mkdir(parents=True, exist_ok=True)
    safe_name = f"{user_id}_{int(time.time())}_{file.filename}"
    file_path = upload_dir / safe_name

    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    document = Document(
        filename=file.filename,
        content_type=file.content_type,
        user_id=user_id,
        status="processing",
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    try:
        text = _extract_text(file_path, file.content_type)
        chunks = _chunk_text(text, settings.CHUNK_SIZE, settings.CHUNK_OVERLAP)
        chunk_records: list[Chunk] = []
        for i, text_chunk in enumerate(chunks):
            chunk = Chunk(document_id=document.id, text=text_chunk, chunk_index=i)
            db.add(chunk)
            chunk_records.append(chunk)
        db.commit()
        for chunk in chunk_records:
            db.refresh(chunk)

        if chunk_records:
            vectors = embedding_engine.embed([c.text for c in chunk_records])
            ids = [c.id for c in chunk_records]
            vector_store.add(ids, vectors)

        document.status = "ready"
        db.commit()
    except Exception as exc:
        document.status = "error"
        document.error_message = str(exc)
        db.commit()
        raise HTTPException(500, f"Failed to process document: {exc}")

    return document


def _extract_text(file_path: Path, content_type: str) -> str:
    if content_type == "application/pdf":
        from pypdf import PdfReader
        reader = PdfReader(str(file_path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        import docx
        doc = docx.Document(str(file_path))
        return "\n".join(p.text for p in doc.paragraphs)
    return file_path.read_text(encoding="utf-8", errors="ignore")


def _chunk_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    text = re.sub(r"\s+", " ", text).strip()
    words = text.split()
    if not words:
        return []
    chunks = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunks.append(" ".join(words[start:end]))
        if end == len(words):
            break
        start = end - chunk_overlap
    return chunks
